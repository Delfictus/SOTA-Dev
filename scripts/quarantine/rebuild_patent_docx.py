"""
Rebuild patent DOCX from clean pdftotext extraction.
Parses section headers, [NNNN] paragraph numbers, claims, and abstract.
Produces a properly formatted DOCX with real Word styles.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_TXT = "/tmp/patent_text.txt"
OUTPUT_DOCX = "/mnt/storage/prism-outputs/docx_export_20260514T080725Z/PRISM4D_provisional_patent_application_CLEAN.docx"

SECTION_HEADERS = {
    "APPLICATION FOR UNITED STATES PROVISIONAL PATENT",
    "TITLE", "CROSS-REFERENCE TO RELATED APPLICATIONS",
    "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH",
    "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH OR DEVELOPMENT",
    "FIELD OF THE INVENTION", "BACKGROUND", "SUMMARY",
    "BRIEF DESCRIPTION OF THE DRAWINGS",
    "DETAILED DESCRIPTION", "DETAILED DESCRIPTION OF THE EMBODIMENTS",
    "CLAIMS", "ABSTRACT", "DRAWINGS", "FIGURES",
    "DESCRIPTION OF EMBODIMENTS", "MODES FOR CARRYING OUT THE INVENTION",
    "INDUSTRIAL APPLICABILITY", "SEQUENCE LISTING",
}

def looks_like_header(line: str) -> bool:
    s = line.strip()
    if not s:
        return False
    # All caps, no period at end, reasonable length
    if s == s.upper() and len(s) > 4 and len(s) < 120 and not s.endswith("."):
        return True
    # Known section names (partial match)
    for h in SECTION_HEADERS:
        if s.startswith(h[:15]):
            return True
    return False

def is_paragraph_tag(line: str) -> bool:
    return bool(re.match(r'^\[\d{4}\]', line.strip()))

def is_claim_line(line: str) -> bool:
    return bool(re.match(r'^(Claim|CLAIM)\s+\d+', line.strip()) or
                re.match(r'^\d+\.\s+[A-Z]', line.strip()))

def set_font(run, size_pt=11, bold=False, color=None):
    run.bold = bold
    run.font.size = Pt(size_pt)
    if color:
        run.font.color.rgb = RGBColor(*color)

def add_section_header(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip())
    set_font(run, size_pt=13, bold=True)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(6)

def add_paragraph_block(doc, tag, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.0)
    p.paragraph_format.first_line_indent = Inches(0.0)
    p.paragraph_format.space_after = Pt(4)
    if tag:
        run_tag = p.add_run(tag + " ")
        set_font(run_tag, size_pt=11, bold=True)
    run_text = p.add_run(text.strip())
    set_font(run_text, size_pt=11)

def add_title_block(doc, lines):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for line in lines:
        run = p.add_run(line.strip() + "\n")
        set_font(run, size_pt=14, bold=True)

def add_meta_line(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text.strip())
    set_font(run, size_pt=11)

# ── Read and parse ──────────────────────────────────────────────────────────
with open(INPUT_TXT, "r", encoding="utf-8", errors="replace") as f:
    raw = f.read()

# Split on form feeds (page breaks)
pages = raw.split("\x0c")
lines = []
for page in pages:
    for line in page.splitlines():
        lines.append(line)

# ── Build document ───────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.25)
    section.right_margin = Inches(1.25)

# Default style
style = doc.styles["Normal"]
style.font.name = "Times New Roman"
style.font.size = Pt(11)

i = 0
title_lines = []
in_title = True
current_para_tag = None
current_para_text = []
in_claims = False

def flush_paragraph(doc, tag, text_parts):
    text = " ".join(t.strip() for t in text_parts if t.strip())
    if text:
        add_paragraph_block(doc, tag, text)

while i < len(lines):
    line = lines[i]
    stripped = line.strip()

    # Skip pure page numbers
    if re.match(r'^\d{1,3}$', stripped):
        i += 1
        continue

    # Empty line — flush current paragraph
    if not stripped:
        if current_para_text:
            flush_paragraph(doc, current_para_tag, current_para_text)
            current_para_tag = None
            current_para_text = []
        i += 1
        continue

    # Title block (first few lines before [0001])
    if in_title and not is_paragraph_tag(stripped) and not looks_like_header(stripped):
        if stripped.startswith("Inventor:") or stripped.startswith("Applicant") or \
           stripped.startswith("Residence") or stripped.startswith("This disclosure"):
            if title_lines:
                add_title_block(doc, title_lines)
                title_lines = []
            add_meta_line(doc, stripped)
        else:
            title_lines.append(stripped)
        i += 1
        continue

    # First paragraph tag ends title mode
    if in_title and is_paragraph_tag(stripped):
        if title_lines:
            add_title_block(doc, title_lines)
            title_lines = []
        in_title = False

    # Section header
    if looks_like_header(stripped) and not is_paragraph_tag(stripped):
        if current_para_text:
            flush_paragraph(doc, current_para_tag, current_para_text)
            current_para_tag = None
            current_para_text = []
        if stripped.upper() in ("CLAIMS", "CLAIM"):
            in_claims = True
        add_section_header(doc, stripped)
        i += 1
        continue

    # Paragraph with [NNNN] tag
    if is_paragraph_tag(stripped):
        if current_para_text:
            flush_paragraph(doc, current_para_tag, current_para_text)
        m = re.match(r'^(\[\d{4}\])\s*(.*)', stripped)
        current_para_tag = m.group(1)
        current_para_text = [m.group(2)] if m.group(2) else []
        i += 1
        continue

    # Claim line
    if in_claims and is_claim_line(stripped):
        if current_para_text:
            flush_paragraph(doc, current_para_tag, current_para_text)
        current_para_tag = None
        current_para_text = [stripped]
        i += 1
        continue

    # Continuation of current paragraph
    if current_para_tag is not None or current_para_text:
        current_para_text.append(stripped)
    else:
        # Standalone line (sub-header, figure label, etc.)
        p = doc.add_paragraph()
        run = p.add_run(stripped)
        set_font(run, size_pt=11)

    i += 1

# Flush any remaining paragraph
if current_para_text:
    flush_paragraph(doc, current_para_tag, current_para_text)

doc.save(OUTPUT_DOCX)
print(f"Saved: {OUTPUT_DOCX}")
print(f"Paragraphs: {len(doc.paragraphs)}")
