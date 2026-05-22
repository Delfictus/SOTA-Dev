#!/usr/bin/env python3
"""
Split each multi-panel structural figure into individual pages.
Portrait figures (Figs 10, 11, 12) are detected, split at their inter-panel
white band, and each half placed on its own full-width page.
All landscape data figures and all text are UNTOUCHED from v3.
"""

import io
import tempfile
import os
import numpy as np
from PIL import Image
from pathlib import Path
from docx import Document
from docx.shared import Emu
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH

INPUT  = Path("/mnt/storage/prism-outputs/docx_export_20260514T080725Z/"
              "PRISM4D_final_structural_integrated_v3.docx")
OUTPUT = Path("/mnt/storage/prism-outputs/docx_export_20260514T080725Z/"
              "PRISM4D_final_structural_integrated_v5_panels_per_page.docx")

EMU              = 914400
CAPTION_RESERVE  = int(0.65 * EMU)


# ── helpers ──────────────────────────────────────────────────────────────────

def find_split_row(img: Image.Image, lo_frac=0.20, hi_frac=0.80,
                   min_brightness=253, min_band=5) -> int:
    """Return the centre row of the widest near-white horizontal band."""
    arr = np.array(img.convert('L'), dtype=np.float32)
    h = arr.shape[0]
    row_bright = arr.mean(axis=1)
    lo, hi = int(h * lo_frac), int(h * hi_frac)
    white = row_bright >= min_brightness
    best_start, best_len = lo, 0
    cur_start, cur_len, in_band = lo, 0, False
    for r in range(lo, hi):
        if white[r]:
            if not in_band:
                cur_start, cur_len, in_band = r, 1, True
            else:
                cur_len += 1
        else:
            if in_band:
                if cur_len > best_len:
                    best_start, best_len = cur_start, cur_len
                in_band = False
    if in_band and cur_len > best_len:
        best_start, best_len = cur_start, cur_len
    if best_len < min_band:
        return h // 2   # fallback: split in half
    return best_start + best_len // 2


def fit_emu(img_w: int, img_h: int, max_w: int, max_h: int):
    """Return (w, h) in EMU fitting inside max_w × max_h, aspect preserved."""
    aspect = img_w / img_h
    page_r = max_w / max_h
    if aspect >= page_r:
        return max_w, int(max_w / aspect)
    else:
        return int(max_h * aspect), max_h


def make_page_break_elem():
    p  = OxmlElement('w:p')
    r  = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    return p


def set_para_pb_center(para):
    para.paragraph_format.page_break_before = True
    para.paragraph_format.alignment         = WD_ALIGN_PARAGRAPH.CENTER


def find_caption(paras, after_idx):
    for j in range(1, 5):
        k = after_idx + j
        if k >= len(paras):
            break
        p = paras[k]
        if p.style.name == 'Caption':
            return k, p
        txt = p.text.strip()
        if txt.lower().startswith('figure') and len(txt) > 6:
            return k, p
        # stop if non-image non-empty paragraph that is not a caption
        has_img = bool(p._element.findall('.//' + qn('a:blip')))
        if txt and not has_img:
            break
    return None, None


# ── main ─────────────────────────────────────────────────────────────────────

def main():
    doc   = Document(INPUT)
    sec   = doc.sections[0]
    uw    = sec.page_width  - sec.left_margin  - sec.right_margin
    uh    = sec.page_height - sec.top_margin   - sec.bottom_margin - CAPTION_RESERVE

    paras = doc.paragraphs
    elem_to_idx = {p._element: i for i, p in enumerate(paras)}

    # Collect portrait shapes (aspect < 1.0)
    portrait_info = []
    for shape in doc.inline_shapes:
        if shape.width / shape.height >= 1.0:
            continue
        blip  = shape._inline.find('.//' + qn('a:blip'))
        embed = blip.get(qn('r:embed'))
        blob  = doc.part.related_parts[embed].blob
        img   = Image.open(io.BytesIO(blob))
        p_elem = shape._inline.getparent().getparent().getparent()
        pidx   = elem_to_idx.get(p_elem)
        portrait_info.append(dict(embed=embed, img=img,
                                  p_elem=p_elem, pidx=pidx))

    print(f"Portrait figures to split: {len(portrait_info)}")
    tmp_files = []

    for info in portrait_info:
        img   = info['img']
        pidx  = info['pidx']
        p_elem = info['p_elem']

        split = find_split_row(img)
        print(f"  Para[{pidx}]: {img.size}px  split at row {split} ({100*split/img.size[1]:.1f}%)")

        top_img = img.crop((0, 0,     img.size[0], split))
        bot_img = img.crop((0, split, img.size[0], img.size[1]))

        # Save halves to temp PNGs
        tf_t = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        tf_b = tempfile.NamedTemporaryFile(suffix='.png', delete=False)
        top_img.save(tf_t.name, 'PNG', dpi=(300, 300))
        bot_img.save(tf_b.name, 'PNG', dpi=(300, 300))
        tmp_files += [tf_t.name, tf_b.name]
        tf_t.close(); tf_b.close()

        top_w, top_h = fit_emu(top_img.width, top_img.height, uw, uh)
        bot_w, bot_h = fit_emu(bot_img.width, bot_img.height, uw, uh)

        # Add both images to doc (appended temporarily to body)
        p_top = doc.add_paragraph()
        p_top.add_run().add_picture(tf_t.name, width=Emu(top_w), height=Emu(top_h))
        set_para_pb_center(p_top)

        p_bot = doc.add_paragraph()
        p_bot.add_run().add_picture(tf_b.name, width=Emu(bot_w), height=Emu(bot_h))
        set_para_pb_center(p_bot)

        top_elem = p_top._element
        bot_elem = p_bot._element

        # Detach them from end of body (will re-attach in correct position)
        body = doc.element.body
        body.remove(top_elem)
        body.remove(bot_elem)

        # Find caption paragraph (uses stale paras list — elements still valid)
        cap_idx, cap_para = find_caption(paras, pidx)

        # Structural replacement:
        # Before: [p_elem(orig image)] [cap_para?] ...
        # After:  [top_elem] [pb] [bot_elem] [cap_para] [pb] ...

        pb_mid   = make_page_break_elem()
        pb_after = make_page_break_elem()

        # Replace original image paragraph with top half
        p_elem.getparent().replace(p_elem, top_elem)

        # Insert pb_mid and bot_elem right after top
        top_elem.addnext(bot_elem)
        top_elem.addnext(pb_mid)

        # Add page break after caption (or after bot if no caption)
        if cap_para:
            cap_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap_para._element.addnext(pb_after)
        else:
            bot_elem.addnext(pb_after)

        print(f"    top: {top_w/EMU:.2f}\" × {top_h/EMU:.2f}\"  "
              f"bot: {bot_w/EMU:.2f}\" × {bot_h/EMU:.2f}\"")

    doc.save(OUTPUT)
    for f in tmp_files:
        try: os.unlink(f)
        except: pass
    print(f"\nSaved → {OUTPUT}")


if __name__ == '__main__':
    main()
