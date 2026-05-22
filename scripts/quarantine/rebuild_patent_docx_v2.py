"""
Rebuild patent DOCX from clean pdftotext extraction — v2 final.
Pre-processes lines to remove page-number artifacts, then standard state machine.
"""
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

INPUT_TXT = "/tmp/patent_text.txt"
OUTPUT_DOCX = "/mnt/storage/prism-outputs/docx_export_20260514T080725Z/PRISM4D_provisional_patent_application_CLEAN.docx"

KNOWN_HEADERS = {
    "APPLICATION FOR UNITED STATES PROVISIONAL PATENT",
    "TITLE",
    "CROSS-REFERENCE TO RELATED APPLICATIONS",
    "STATEMENT REGARDING FEDERALLY SPONSORED RESEARCH",
    "OR DEVELOPMENT",
    "FIELD OF THE INVENTION",
    "BACKGROUND",
    "SUMMARY",
    "DEFINITIONS",
    "BRIEF DESCRIPTION OF THE DRAWINGS",
    "DETAILED DESCRIPTION",
    "EXEMPLARY CLAIMS",
    "ABSTRACT",
    "DRAWINGS",
}

def is_page_num(s):
    return bool(re.match(r'^\d{1,3}$', s))

def is_para_tag(s):
    return bool(re.match(r'^\[\d{4}\]', s))

def is_section_header(s):
    return s.strip() in KNOWN_HEADERS

def is_claim_start(s):
    return bool(re.match(r'^\d{1,2}\.\s+[A-Z]', s))

def keep_with_next(para):
    pPr = para._p.get_or_add_pPr()
    kwn = OxmlElement('w:keepWithNext')
    pPr.append(kwn)

def add_page_number_footer(doc):
    for section in doc.sections:
        footer = section.footer
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.clear()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run()
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        for tag, instr in [('begin', None), (None, ' PAGE '), ('end', None)]:
            if tag:
                e = OxmlElement('w:fldChar')
                e.set(qn('w:fldCharType'), tag)
                run._r.append(e)
            else:
                e = OxmlElement('w:instrText')
                e.text = instr
                run._r.append(e)

# ── Read and pre-process ──────────────────────────────────────────────────────
with open(INPUT_TXT, "r", encoding="utf-8", errors="replace") as f:
    raw = f.read()
raw = raw.replace("\x0c", "\n")
raw_lines = [l.strip() for l in raw.splitlines()]

def is_new_para_start(s):
    """True if this line clearly starts a new paragraph, not a continuation."""
    if not s:
        return False
    if is_para_tag(s) or is_section_header(s) or is_claim_start(s):
        return True
    # Roman numeral sub-section headers: I., II., III. ...
    if re.match(r'^[IVX]+\.\s+[A-Z]', s):
        return True
    return False

# Pre-processing: remove ALL page-break artifacts.
# Handles both:
#   A) content \n page_num \n blank \n continuation  (page num right after content)
#   B) content \n blank \n page_num \n blank \n continuation  (blank before page num)
# In both cases: remove page_num and surrounding blanks.
# Exception: if what follows the page break is a new paragraph start, keep one blank.
lines = raw_lines
cleaned = []
j = 0
while j < len(lines):
    line = lines[j]

    # Case: blank line — check if page number follows (pattern B)
    if not line:
        # Scan ahead past blanks to find next non-blank
        k = j + 1
        while k < len(lines) and not lines[k]:
            k += 1
        if k < len(lines) and is_page_num(lines[k]):
            # Blank-before-page-number pattern: skip the blank and the page number
            j = k + 1  # position after page number
            # Skip any blanks after the page number too
            while j < len(lines) and not lines[j]:
                j += 1
            # What comes next?
            next_content = lines[j] if j < len(lines) else ''
            if is_new_para_start(next_content):
                cleaned.append('')  # Keep one blank (real paragraph boundary)
            # else: mid-paragraph, emit nothing, continue from j
            continue
        # Normal blank line (no page number follows)
        cleaned.append(line)
        j += 1
        continue

    # Case: page number line (no blank before it in cleaned stream)
    if is_page_num(line):
        j += 1
        # Skip any blanks after the page number
        while j < len(lines) and not lines[j]:
            j += 1
        next_content = lines[j] if j < len(lines) else ''
        if is_new_para_start(next_content):
            cleaned.append('')  # Keep one blank (real paragraph boundary)
        # else: mid-paragraph, no blank needed
        continue

    # Regular line
    cleaned.append(line)
    j += 1

all_lines = cleaned
n = len(all_lines)

# Find title block end
title_end = next((i for i, l in enumerate(all_lines) if is_section_header(l)), 0)

# ── Build semantic blocks ─────────────────────────────────────────────────────
blocks = []
cur_tag = None
cur_parts = []
pending_header = None
in_claims = False
in_abstract = False
in_drawings = False

def emit_pending():
    global cur_tag, cur_parts
    if cur_tag is not None or cur_parts:
        text = ' '.join(p for p in cur_parts if p)
        blocks.append(('para', cur_tag, text))
    cur_tag = None
    cur_parts = []

def emit_header(h):
    global in_claims, in_abstract, in_drawings
    emit_pending()
    blocks.append(('section_header', h))
    u = h.upper()
    if 'CLAIM' in u:
        in_claims = True; in_abstract = False; in_drawings = False
    elif u.strip() == 'ABSTRACT':
        in_abstract = True; in_claims = False; in_drawings = False
    elif u.strip() == 'DRAWINGS':
        in_drawings = True; in_claims = False; in_abstract = False
    else:
        in_claims = False; in_abstract = False

# Title block
i = 0
while i < title_end:
    line = all_lines[i]
    if not line:
        i += 1; continue
    if re.match(r'^DELF-', line):
        blocks.append(('title_ref', line))
    elif line.startswith('Provisional Patent'):
        blocks.append(('title_sub', line))
    elif re.match(r'^(Inventor|Applicant|Residence)', line):
        blocks.append(('title_meta', line))
    elif line.startswith('This disclosure') or (blocks and blocks[-1][0] == 'title_blurb'):
        if blocks and blocks[-1][0] == 'title_blurb':
            blocks[-1] = ('title_blurb', blocks[-1][1] + ' ' + line)
        else:
            blocks.append(('title_blurb', line))
    else:
        if blocks and blocks[-1][0] == 'title_main':
            blocks[-1] = ('title_main', blocks[-1][1] + ' ' + line)
        else:
            blocks.append(('title_main', line))
    i += 1

# Main body
i = title_end
while i < n:
    line = all_lines[i]
    i += 1

    # Blank line → flush current paragraph (genuine paragraph boundary)
    if not line:
        if pending_header is not None:
            emit_header(pending_header)
            pending_header = None
        else:
            emit_pending()
        continue

    # Section header
    if is_section_header(line):
        s = line.strip()
        if s == 'OR DEVELOPMENT':
            if pending_header is not None:
                pending_header = pending_header + ' ' + s
            else:
                pending_header = s
        else:
            if pending_header is not None:
                emit_header(pending_header)
            pending_header = s
        continue

    # Any body content: emit pending header first
    if pending_header is not None:
        emit_header(pending_header)
        pending_header = None

    # [NNNN] paragraph tag
    if is_para_tag(line):
        emit_pending()
        m = re.match(r'^(\[\d{4}\])\s*(.*)', line)
        cur_tag = m.group(1)
        rest = m.group(2).strip()
        cur_parts = [rest] if rest else []
        continue

    # Claim
    if in_claims and is_claim_start(line):
        emit_pending()
        cur_tag = None
        cur_parts = [line]
        continue

    # Abstract: one big paragraph
    if in_abstract:
        if cur_parts:
            cur_parts.append(line)
        else:
            cur_tag = None
            cur_parts = [line]
        continue

    # Drawings: only keep FIG. labels + the brief intro paragraph
    if in_drawings:
        if re.match(r'^FIG\.\s*\d', line):
            emit_pending()
            blocks.append(('fig_label', line))
        elif blocks and blocks[-1][0] == 'section_header' and blocks[-1][1] == 'DRAWINGS':
            blocks.append(('drawings_intro', line))
        elif blocks and blocks[-1][0] == 'drawings_intro':
            blocks[-1] = ('drawings_intro', blocks[-1][1] + ' ' + line)
        # All diagram node labels: silently drop
        continue

    # Normal body continuation
    if cur_tag is not None or cur_parts:
        cur_parts.append(line)
    else:
        # Standalone line (sub-section like "I. Platform Overview")
        blocks.append(('subhead', line))

# Final flush
if pending_header:
    emit_header(pending_header)
emit_pending()

# ── Post-process: merge consecutive subhead blocks ────────────────────────────
# Consecutive subheads are table cells or split text — join them into one block.
# A lone subhead that matches Roman numeral pattern becomes a 'subsection'.
ROMAN_SUBHEAD = re.compile(r'^[IVX]+\.\s+[A-Z]')
merged = []
sbuf = []
for block in blocks:
    if block[0] == 'subhead':
        sbuf.append(block[1])
    else:
        if sbuf:
            if len(sbuf) == 1 and ROMAN_SUBHEAD.match(sbuf[0]):
                merged.append(('subsection', sbuf[0]))
            else:
                # Merge all cells/fragments into one body paragraph
                # Join with space; use dash separator for obvious table rows
                merged.append(('table_text', ' '.join(sbuf)))
            sbuf = []
        merged.append(block)
if sbuf:
    if len(sbuf) == 1 and ROMAN_SUBHEAD.match(sbuf[0]):
        merged.append(('subsection', sbuf[0]))
    else:
        merged.append(('table_text', ' '.join(sbuf)))
blocks = merged

# ── Build Word document ───────────────────────────────────────────────────────
doc = Document()
for sec in doc.sections:
    sec.top_margin = Inches(1.0)
    sec.bottom_margin = Inches(1.0)
    sec.left_margin = Inches(1.25)
    sec.right_margin = Inches(1.25)

doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(11)
doc.styles['Normal'].paragraph_format.space_after = Pt(6)

h1 = doc.styles['Heading 1']
h1.font.name = 'Times New Roman'
h1.font.size = Pt(12)
h1.font.bold = True
h1.font.color.rgb = None
h1.paragraph_format.space_before = Pt(18)
h1.paragraph_format.space_after = Pt(6)
h1.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
h1.paragraph_format.keep_with_next = True

add_page_number_footer(doc)
CENTER = WD_ALIGN_PARAGRAPH.CENTER

def txt(text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT, sb=0, sa=6):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after = Pt(sa)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    run.font.bold = bold
    return p

def heading(text):
    p = doc.add_paragraph(style='Heading 1')
    p.clear()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.font.bold = True
    keep_with_next(p)
    return p

def para_block(tag, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    if tag:
        rt = p.add_run(tag + ' ')
        rt.font.name = 'Times New Roman'
        rt.font.size = Pt(11)
        rt.bold = True
    rb = p.add_run(text)
    rb.font.name = 'Times New Roman'
    rb.font.size = Pt(11)
    return p

for block in blocks:
    btype = block[0]
    if btype == 'title_ref':
        txt(block[1], size=11, align=CENTER, sb=12, sa=4)
    elif btype == 'title_main':
        txt(block[1], bold=True, size=14, align=CENTER, sb=2, sa=4)
    elif btype == 'title_sub':
        txt(block[1], size=11, align=CENTER, sb=8, sa=4)
    elif btype == 'title_meta':
        txt(block[1], size=11, align=CENTER, sb=2, sa=2)
    elif btype == 'title_blurb':
        txt(block[1], size=11, align=CENTER, sb=12, sa=18)
    elif btype == 'section_header':
        heading(block[1])
    elif btype == 'para':
        _, tag, text = block
        para_block(tag, text)
    elif btype == 'subsection':
        p = txt(block[1], bold=True, size=11, sb=10, sa=4)
        keep_with_next(p)
    elif btype == 'table_text':
        txt(block[1], size=10, sb=2, sa=4)
    elif btype == 'drawings_intro':
        txt(block[1], size=11)
    elif btype == 'fig_label':
        p = txt(block[1], bold=True, size=11, sb=10, sa=4)
        keep_with_next(p)

doc.save(OUTPUT_DOCX)
print(f"Saved: {OUTPUT_DOCX}")
print(f"Paragraphs: {len(doc.paragraphs)}")
from collections import Counter
counts = Counter(b[0] for b in blocks)
for k, v in sorted(counts.items()):
    print(f"  {k}: {v}")
